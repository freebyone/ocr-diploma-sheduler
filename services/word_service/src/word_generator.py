# word_generator.py — полная финальная версия

import os
import re
from collections import defaultdict
from typing import List, Optional, Tuple, Dict

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml

from sqlalchemy.orm import Session, joinedload

from models import (
    IncomingDirection, Student, Specialization,
    ControlTable, StudyProgram, FormatControl, FormatRetests,
)

# ══════════════════════════════════════════════
#  УТИЛИТЫ ФОРМАТИРОВАНИЯ
# ══════════════════════════════════════════════

def fix_font(run) -> None:
    """
    Принудительно прописывает Times New Roman в XML,
    чтобы тема документа не подменяла шрифт.
    """
    run.font.name = 'Times New Roman'
    # rPr может ещё не существовать — создаём безопасно
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'),    'Times New Roman')
    rFonts.set(qn('w:hAnsi'),   'Times New Roman')
    rFonts.set(qn('w:eastAsia'),'Times New Roman')
    rFonts.set(qn('w:cs'),      'Times New Roman')


def set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    size: int = 10,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    """Записать текст в ячейку с форматированием."""
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    run = para.add_run(str(text) if text else '')
    run.font.size = Pt(size)
    run.bold = bold
    fix_font(run)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_shading(cell, color: str = "FFFFFF") -> None:
    """Заливка ячейки (по умолчанию белая)."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_paragraph_text(
    doc: Document,
    text: str,
    bold: bool = False,
    size: int = 12,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    space_after: int = 6,
    space_before: int = 0,
    first_line_indent: Optional[float] = None,
):
    """Добавить абзац с форматированием."""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.space_before = Pt(space_before)

    if first_line_indent is not None:
        para.paragraph_format.first_line_indent = Cm(first_line_indent)

    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    fix_font(run)
    return para

# ══════════════════════════════════════════════
#  ТАБЛИЦА ПЕРЕАТТЕСТАЦИИ
# ══════════════════════════════════════════════

# Ширины колонок в сантиметрах (сумма ≈ 16.5 см под A4 с полями 3+1.5)
_COL_WIDTHS_CM = [5.0, 1.5, 2.2, 1.5, 2.2, 1.6, 2.5]

def add_control_table(doc: Document, control_rows: List[ControlTable]):
    """
    Строит таблицу переаттестации:

    | Дисциплина | По уч. плану      | Данные диплома    | Часов | Форма |
    |            | К-во ч | Форма    | К-во ч | Форма    | пере- | пере- |
    |            |        | контроля |        | контроля | аттест| аттест|
    """
    total_rows = 2 + len(control_rows)
    total_cols = 7

    table = doc.add_table(rows=total_rows, cols=total_cols)
    table.style     = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit   = False  # фиксированная ширина

    # Применяем ширины ко всем ячейкам каждой колонки
    for col_idx, w in enumerate(_COL_WIDTHS_CM):
        for cell in table.columns[col_idx].cells:
            cell.width = Cm(w)

    # ── Строка 0: крупные заголовки ──────────────────────────────────────

    # Колонка 0: «Наименование дисциплины» — объединяем строки 0-1
    cell00 = table.cell(0, 0)
    cell00.merge(table.cell(1, 0))
    set_cell_text(
        cell00, 'Наименование\nдисциплины',
        bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_shading(cell00)

    # Колонки 1-2: «По учебному плану» — объединяем по горизонтали
    cell01 = table.cell(0, 1)
    cell01.merge(table.cell(0, 2))
    set_cell_text(
        cell01, 'По учебному плану',
        bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_shading(cell01)

    # Колонки 3-4: «Данные приложений к диплому»
    cell03 = table.cell(0, 3)
    cell03.merge(table.cell(0, 4))
    set_cell_text(
        cell03, 'Данные приложений\nк диплому',
        bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_shading(cell03)

    # Колонка 5: «Переаттестовано (часов)» — объединяем строки 0-1
    cell05 = table.cell(0, 5)
    cell05.merge(table.cell(1, 5))
    set_cell_text(
        cell05, 'Пере-\nаттестовано\n(часов)',
        bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_shading(cell05)

    # Колонка 6: «Форма переаттестации» — объединяем строки 0-1
    cell06 = table.cell(0, 6)
    cell06.merge(table.cell(1, 6))
    set_cell_text(
        cell06, 'Форма\nпереаттестации',
        bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    set_cell_shading(cell06)

    # ── Строка 1: подзаголовки ───────────────────────────────────────────
    sub = [
        (1, 'К-во\nчасов'),
        (2, 'Форма\nконтроля'),
        (3, 'К-во\nчасов'),
        (4, 'Форма\nконтроля'),
    ]
    for col_idx, text in sub:
        c = table.cell(1, col_idx)
        set_cell_text(c, text, bold=True, size=8,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(c)

    # ── Строки данных ────────────────────────────────────────────────────
    for i, ct in enumerate(control_rows):
        r = i + 2

        prog_name     = ct.study_program.name       if ct.study_program       else ''
        fc_norma_name = ct.format_control_norma.format_name \
                        if ct.format_control_norma else ''
        fc_fact_name  = ct.format_control_fact.format_name  \
                        if ct.format_control_fact  else ''
        fr_name       = ct.format_retests.format_name       \
                        if ct.format_retests       else ''

        set_cell_text(table.cell(r, 0), prog_name,             size=9)
        set_cell_text(table.cell(r, 1), ct.hours_normal or '',  size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(r, 2), fc_norma_name,          size=8)
        set_cell_text(table.cell(r, 3), ct.hours_fact or '',    size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(r, 4), fc_fact_name,           size=8)
        set_cell_text(table.cell(r, 5), ct.hours_fact or '',    size=9,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(table.cell(r, 6), fr_name,                size=8)

    return table

# ══════════════════════════════════════════════
#  ГРУППИРОВКА СТУДЕНТОВ
# ══════════════════════════════════════════════

GroupKey = Tuple[int, str, int, str]  # (uni_id, uni_name, spec_id, spec_name)


def group_students_by_uni_spec(
    students: List[Student],
) -> Dict[GroupKey, List[Student]]:
    """
    Группирует студентов по (University, Specialization).
    Внутри каждой группы сортирует по ФИО.
    """
    groups: Dict[GroupKey, List[Student]] = defaultdict(list)

    for student in students:
        spec = student.specialization
        if spec and spec.university:
            key: GroupKey = (
                spec.university.id,
                spec.university.name,
                spec.id,
                spec.name,
            )
        elif spec:
            key = (0, 'Учебное заведение не указано', spec.id, spec.name)
        else:
            key = (0, 'Учебное заведение не указано',
                   0, 'Специальность не указана')

        groups[key].append(student)

    for key in groups:
        groups[key].sort(key=lambda s: s.full_name)

    return dict(groups)


def get_unique_control_data(
    session: Session,
    direction_id: int,
) -> List[ControlTable]:
    """
    Возвращает уникальные строки ControlTable для направления.
    Дедупликация по study_program_id.
    """
    rows = (
        session.query(ControlTable)
        .filter(ControlTable.incoming_direction_id == direction_id)
        .options(
            joinedload(ControlTable.study_program),
            joinedload(ControlTable.format_control_norma),
            joinedload(ControlTable.format_control_fact),
            joinedload(ControlTable.format_retests),
        )
        .all()
    )

    seen: set = set()
    unique: List[ControlTable] = []
    for ct in rows:
        if ct.study_program_id not in seen:
            seen.add(ct.study_program_id)
            unique.append(ct)

    return unique

# ══════════════════════════════════════════════
#  ГЕНЕРАЦИЯ ОДНОГО ПРИКАЗА
# ══════════════════════════════════════════════

def generate_order_for_direction(
    session: Session,
    direction: IncomingDirection,
    output_dir: str,
) -> str:
    """Генерирует Word-приказ для одного IncomingDirection."""

    # ── Загрузка данных ──────────────────────────────────────────────────
    students = (
        session.query(Student)
        .filter(Student.incoming_direction_id == direction.id)
        .options(
            joinedload(Student.specialization)
            .joinedload(Specialization.university)
        )
        .all()
    )

    if not students:
        raise ValueError(
            f"Нет студентов для направления id={direction.id} «{direction.name}»"
        )

    control_data = get_unique_control_data(session, direction.id)

    if not control_data:
        raise ValueError(
            f"Нет записей ControlTable для направления id={direction.id}"
        )

    groups = group_students_by_uni_spec(students)

    # ── Создание документа ───────────────────────────────────────────────
    doc = Document()

    # Язык документа (ru-RU) — важно для переносов и орфографии
    try:
        styles_element = doc.styles.element
        rpr_nodes = styles_element.xpath(
            './w:docDefaults/w:rPrDefault/w:rPr',
            namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'},
        )
        if rpr_nodes:
            lang = parse_xml(
                f'<w:lang {nsdecls("w")} '
                f'w:val="ru-RU" w:eastAsia="ru-RU" w:bidi="ar-SA"/>'
            )
            rpr_nodes[0].append(lang)
    except Exception:
        pass  # Не критично

    # Поля страницы (A4)
    section = doc.sections[0]
    section.page_height   = Cm(29.7)
    section.page_width    = Cm(21.0)
    section.left_margin   = Cm(3.0)   # под подшивку
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Базовый стиль
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    # ── Содержимое документа ─────────────────────────────────────────────

    add_paragraph_text(
        doc, 'ПРИКАЗ',
        bold=True, size=14,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    add_paragraph_text(
        doc, 'ПРИКАЗЫВАЮ:',
        bold=True, size=12,
        space_after=6,
    )

    main_text = (
        'Переаттестовать дисциплины в соответствии с учебным планом '
        'по ускоренной образовательной программе бакалавриата '
        '(на базе профильного среднего профессионального образования) '
        f'по направлению подготовки {direction.name} '
        'у нижеследующих студентов 1 курса, очно-заочной (вечерней) '
        'формы обучения:'
    )
    add_paragraph_text(
        doc, main_text,
        size=12,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line_indent=1.25,
        space_after=10,
    )

    # ── Блоки по группам ─────────────────────────────────────────────────
    for (uni_id, uni_name, spec_id, spec_name), group_students in groups.items():

        # Список студентов
        for idx, student in enumerate(group_students, 1):
            add_paragraph_text(
                doc,
                f'{idx}. {student.full_name}',
                size=12,
                space_after=2,
            )

        # Связующий текст
        add_paragraph_text(
            doc,
            f'прослушанных в {uni_name} по специальности «{spec_name}»:',
            size=12,
            space_before=6,
            space_after=4,
        )

        # Таблица переаттестации
        add_control_table(doc, control_data)

        # Небольшой отступ после таблицы
        add_paragraph_text(doc, '', size=6, space_after=10)

    # ── Сохранение ───────────────────────────────────────────────────────
    os.makedirs(output_dir, exist_ok=True)

    safe_name = re.sub(r'[^\w\sа-яА-ЯёЁ\-]', '', direction.name)
    safe_name = safe_name.strip().replace(' ', '_')[:60]
    filepath  = os.path.join(output_dir, f'Приказ_{safe_name}.docx')

    doc.save(filepath)
    return filepath