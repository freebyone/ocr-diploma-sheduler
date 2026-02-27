from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Dict, Any
import logging
from sqlalchemy.orm import Session
import io
from docx import Document

# Меняем на абсолютные импорты
from config import get_settings
from database import get_db, init_db
from parser import WordParser

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

settings = get_settings()

# Создаем приложение FastAPI
app = FastAPI(
    title=settings.APP_NAME,
    version=settings.APP_VERSION,
    debug=settings.DEBUG
)

# Настройка CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def startup_event():
    """Действия при запуске приложения"""
    logger.info("Starting up Word Parser Service")
    try:
        init_db()
        logger.info("Database initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize database: {e}")
        raise


@app.get("/")
async def root():
    """Корневой эндпоинт для проверки работы сервиса"""
    return {
        "service": settings.APP_NAME,
        "version": settings.APP_VERSION,
        "status": "running"
    }


@app.get("/health")
async def health_check(db: Session = Depends(get_db)):
    """Проверка здоровья сервиса"""
    try:
        db.execute("SELECT 1")
        return {"status": "healthy", "database": "connected"}
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        raise HTTPException(status_code=500, detail="Database connection failed")


@app.post("/api/parse-word")
async def parse_word_file(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
) -> Dict[str, Any]:
    """Эндпоинт для парсинга одного Word файла"""
    logger.info(f"Received file: {file.filename}")
    
    if not file.filename.endswith(('.docx', '.doc')):
        raise HTTPException(
            status_code=400,
            detail="File must be a Word document (.docx or .doc)"
        )
    
    try:
        content = await file.read()
        doc = Document(io.BytesIO(content))
        
        text_lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_lines.append(para.text)
        
        logger.info(f"Extracted {len(text_lines)} lines from document")
        
        parser = WordParser(db)
        stats = parser.parse_document(text_lines)
        
        return {
            "success": True,
            "filename": file.filename,
            "stats": stats,
            "message": f"Successfully processed document. Created {stats['control_tables']} records."
        }
        
    except Exception as e:
        logger.error(f"Error processing file {file.filename}: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")


@app.post("/api/parse-multiple-word")
async def parse_multiple_word_files(
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db)
) -> Dict[str, Any]:
    """Эндпоинт для парсинга нескольких Word файлов"""
    logger.info(f"Received {len(files)} files")
    
    results = []
    total_stats = {
        'universities': 0,
        'directions': 0,
        'specializations': 0,
        'study_programs': 0,
        'control_tables': 0,
        'errors': []
    }
    
    for file in files:
        file_result = {
            "filename": file.filename,
            "success": False,
            "stats": None,
            "error": None
        }
        
        try:
            if not file.filename.endswith(('.docx', '.doc')):
                file_result["error"] = "Invalid file format"
                results.append(file_result)
                continue
            
            content = await file.read()
            doc = Document(io.BytesIO(content))
            
            text_lines = [para.text for para in doc.paragraphs if para.text.strip()]
            
            parser = WordParser(db)
            stats = parser.parse_document(text_lines)
            
            file_result["success"] = True
            file_result["stats"] = stats
            
            for key in ['universities', 'directions', 'specializations', 'study_programs', 'control_tables']:
                total_stats[key] += stats.get(key, 0)
            total_stats['errors'].extend(stats.get('errors', []))
            
        except Exception as e:
            logger.error(f"Error processing {file.filename}: {e}")
            file_result["error"] = str(e)
            total_stats['errors'].append(f"{file.filename}: {str(e)}")
        
        results.append(file_result)
    
    return {
        "success": True,
        "total_files": len(files),
        "successful": sum(1 for r in results if r["success"]),
        "failed": sum(1 for r in results if not r["success"]),
        "total_stats": total_stats,
        "results": results
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)